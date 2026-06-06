"""
Export utilities: resume/cover letter → DOCX or PDF bytes.
Used by Streamlit's st.download_button.
"""

import io
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _add_section(doc, heading: str, body_lines: list[str]):
    h = doc.add_paragraph(heading)
    h.style = "Heading 1"
    h.runs[0].font.color.rgb = RGBColor(0x23, 0x52, 0x89)
    for line in body_lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("•") or line.startswith("-"):
            p = doc.add_paragraph(line.lstrip("•- ").strip(), style="List Bullet")
        else:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(11) if p.runs else None


def resume_to_docx(resume_text: str, candidate_name: str = "Candidate") -> bytes:
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph(candidate_name)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x23, 0x52, 0x89)

    doc.add_paragraph()  # spacer

    # Parse sections by ALL-CAPS headings
    current_heading = None
    current_lines = []

    for raw_line in resume_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        # Detect ALL-CAPS section headers (≥3 chars, mostly uppercase)
        if line.isupper() and len(line) >= 3:
            if current_heading:
                _add_section(doc, current_heading, current_lines)
            current_heading = line
            current_lines = []
        else:
            current_lines.append(line)

    # Flush last section
    if current_heading:
        _add_section(doc, current_heading, current_lines)
    elif current_lines:
        for line in current_lines:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def cover_letter_to_docx(letter_text: str, candidate_name: str = "Candidate") -> bytes:
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header
    title = doc.add_paragraph("Cover Letter")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x23, 0x52, 0x89)

    sub = doc.add_paragraph(candidate_name)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Body paragraphs
    for para in letter_text.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para.replace("\n", " "))
            p.paragraph_format.space_after = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def interview_to_docx(interview_data: dict, candidate_name: str = "Candidate") -> bytes:
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph("Interview Preparation Guide")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x23, 0x52, 0x89)

    sub = doc.add_paragraph(candidate_name)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    questions = interview_data.get("questions", [])
    for i, q in enumerate(questions, 1):
        # Question number + category
        cat_para = doc.add_paragraph()
        cat_run = cat_para.add_run(f"Q{i}  [{q.get('category', '')}]")
        cat_run.bold = True
        cat_run.font.color.rgb = RGBColor(0x23, 0x52, 0x89)
        cat_run.font.size = Pt(12)

        # Question text
        q_para = doc.add_paragraph(q.get("question", ""))
        q_para.runs[0].bold = True
        q_para.runs[0].font.size = Pt(11)

        # Why asked
        why = q.get("why_asked", "")
        if why:
            w = doc.add_paragraph(f"💡 Why asked: {why}")
            w.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
            w.runs[0].font.size = Pt(10)
            w.runs[0].italic = True

        # Suggested answer
        ans_label = doc.add_paragraph("Suggested Answer:")
        ans_label.runs[0].bold = True
        ans_label.runs[0].font.size = Pt(10)

        ans = doc.add_paragraph(q.get("suggested_answer", ""))
        ans.runs[0].font.size = Pt(11)
        ans.paragraph_format.left_indent = Inches(0.25)
        ans.paragraph_format.space_after = Pt(16)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

ACCENT = colors.HexColor("#235289")
MUTED  = colors.HexColor("#595959")
BODY   = colors.HexColor("#1a1a1a")


def _base_styles():
    styles = getSampleStyleSheet()
    return {
        "name": ParagraphStyle("Name", fontSize=18, textColor=ACCENT,
                               alignment=TA_CENTER, spaceAfter=4, fontName="Helvetica-Bold"),
        "section": ParagraphStyle("Section", fontSize=11, textColor=ACCENT,
                                  fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=4,
                                  textTransform="uppercase"),
        "body": ParagraphStyle("Body", fontSize=10, textColor=BODY,
                               fontName="Helvetica", leading=15, spaceAfter=4,
                               alignment=TA_JUSTIFY),
        "bullet": ParagraphStyle("Bullet", fontSize=10, textColor=BODY,
                                 fontName="Helvetica", leading=15, spaceAfter=3,
                                 leftIndent=14, bulletIndent=0),
        "title": ParagraphStyle("Title", fontSize=16, textColor=ACCENT,
                                alignment=TA_CENTER, spaceAfter=2, fontName="Helvetica-Bold"),
        "subtitle": ParagraphStyle("Subtitle", fontSize=11, textColor=MUTED,
                                   alignment=TA_CENTER, spaceAfter=12, fontName="Helvetica"),
        "para": ParagraphStyle("Para", fontSize=10.5, textColor=BODY,
                               fontName="Helvetica", leading=16, spaceAfter=10,
                               alignment=TA_JUSTIFY),
    }


def resume_to_pdf(resume_text: str, candidate_name: str = "Candidate") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            topMargin=0.75*inch, bottomMargin=0.75*inch,
                            leftMargin=inch, rightMargin=inch)
    s = _base_styles()
    story = []

    story.append(Paragraph(candidate_name, s["name"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT, spaceAfter=8))

    current_heading = None
    current_lines = []

    def flush():
        if current_heading:
            story.append(Spacer(1, 4))
            story.append(Paragraph(current_heading, s["section"]))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))
            for line in current_lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("•") or line.startswith("-"):
                    clean = line.lstrip("•- ").strip()
                    story.append(Paragraph(f"• {clean}", s["bullet"]))
                else:
                    story.append(Paragraph(line, s["body"]))

    for raw_line in resume_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.isupper() and len(line) >= 3:
            flush()
            current_heading = line
            current_lines = []
        else:
            current_lines.append(line)
    flush()

    doc.build(story)
    return buf.getvalue()


def cover_letter_to_pdf(letter_text: str, candidate_name: str = "Candidate") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            topMargin=inch, bottomMargin=inch,
                            leftMargin=1.25*inch, rightMargin=1.25*inch)
    s = _base_styles()
    story = []

    story.append(Paragraph("Cover Letter", s["title"]))
    story.append(Paragraph(candidate_name, s["subtitle"]))
    story.append(HRFlowable(width="100%", thickness=1, color=ACCENT, spaceAfter=16))

    for para in letter_text.split("\n\n"):
        para = para.strip().replace("\n", " ")
        if para:
            story.append(Paragraph(para, s["para"]))

    doc.build(story)
    return buf.getvalue()
